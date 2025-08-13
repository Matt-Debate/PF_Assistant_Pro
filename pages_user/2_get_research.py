import streamlit as st
import requests
import httpx
from datetime import datetime
from auth import get_username, log_query, get_remaining_queries
from docx import Document
from io import BytesIO
from requests.exceptions import SSLError, RequestException

# --- Constants ---
TOOL_NAME = "get_research"
MAX_DAILY_QUERIES = 20
PERPLEXITY_API_KEY = st.secrets["PERPLEXITY_API_KEY"]
PERPLEXITY_ENDPOINT = "https://api.perplexity.ai/chat/completions"

# --- Prompt Builder ---
def build_prompt(topic, side, claim, quant):
    if topic and side:
        preface = f"For the debate topic '{topic}' on the {side} side"
    elif topic:
        preface = f"For the debate topic '{topic}'"
    else:
        preface = ""

    prompt = f"{preface + ', ' if preface else ''}produce 3 evidence cards that support this claim: {claim}."

    prompt += (
        " 1 evidence card can only have 1 quote and 1 source. "
        "Before outputting to the user, verify that the cards match the claim. "
        "If no matching quotes are found, respond 'no cards found'. "
        "If only 1 or 2 quotes are found, only provide the matching quotes. "
    )

    if quant:
        prompt += "Find quotes with statistical data. "

    prompt += (
        "Format cards exactly like this, and provide no additional commentary or markdown:\n\n"
        "Author(or publishing organization if no author is found), Year of Publication\n"
        "URL\n"
        "\"Direct Quote\""
    )

    return prompt


# --- API Query ---
def query_perplexity(prompt):
    headers = {
        "Authorization": f"Bearer {PERPLEXITY_API_KEY}",
        "Accept": "application/json",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "sonar",
        "messages": [
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.5
    }

    try:
        response = requests.post(PERPLEXITY_ENDPOINT, headers=headers, json=payload, timeout=30)
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except SSLError:
        with httpx.Client(http2=True, verify=True, timeout=30) as client:
            response = client.post(PERPLEXITY_ENDPOINT, headers=headers, json=payload)
            response.raise_for_status()
            return response.json()["choices"][0]["message"]["content"]
    except RequestException as e:
        raise RuntimeError(f"Request failed: {e}") from e


# --- DOCX Output ---
def create_docx_output(text):
    if not isinstance(text, str):
        return None
    doc = Document()
    doc.add_heading("Evidence Cards", level=1)
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- Initialize Session State ---
if "research_input" not in st.session_state:
    st.session_state.research_input = {}
if "evidence_result" not in st.session_state:
    st.session_state.evidence_result = None

# --- Title ---
st.title("📚 Evidence Machine")
st.subheader("Let’s Cut Some Cards!")

# --- Auth & Limits ---
username = get_username()
remaining = get_remaining_queries(username) or MAX_DAILY_QUERIES

if remaining <= 0:
    st.warning("You have no uses remaining today.")
    st.stop()


# --- Inputs ---
if st.session_state.evidence_result is None:
    with st.form("research_form"):
        topic = st.text_input("Topic (optional)", value=st.session_state.research_input.get("topic", ""))
        side = st.selectbox("Side (optional)", ["", "Pro", "Con"], 
                            index=["", "Pro", "Con"].index(st.session_state.research_input.get("side", "")))
        claim = st.text_area("Claim to find evidence for (required)", 
                             height=120, value=st.session_state.research_input.get("claim", ""))
        quant = st.checkbox("Quantitative or statistical data", value=st.session_state.research_input.get("quant", False))
        submitted = st.form_submit_button("Submit")

    st.markdown(f"**You have {remaining} uses of {TOOL_NAME} remaining today.**")
    st.info("⚠️ Save your result before searching again. New results will replace old ones.")

    if submitted:
        if not claim.strip():
            st.error("Please enter a claim.")
            st.stop()

        # Save inputs
        st.session_state.research_input = {
            "topic": topic,
            "side": side,
            "claim": claim,
            "quant": quant
        }

        # Run search
        with st.spinner("Cutting your card(s)..."):
            try:
                prompt = build_prompt(topic, side.lower() if side else "", claim, quant)
                raw_result = query_perplexity(prompt)
                st.session_state.evidence_result = raw_result
                log_query(username, TOOL_NAME, True)
                st.success("✅ Evidence generated!")
                st.rerun()
            except Exception as e:
                st.error(f"❌ Error fetching evidence: {e}")
                log_query(username, TOOL_NAME, False)


# --- Outputs ---
else:
    raw_result = st.session_state.evidence_result

    # Dynamically set height based on number of lines (min 300px)
    line_count = raw_result.count("\n") + 1
    text_height = max(300, line_count * 20)  # 20px per line approx.

    with st.expander("🔍 Evidence Output", expanded=True):
        st.text_area("Cards", raw_result, height=text_height)

    # --- Buttons ---
    docx_buffer = create_docx_output(raw_result)
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        if docx_buffer:
            st.download_button("📥 Download", docx_buffer, file_name="evidence.docx")
    with col2:
        if st.button("🔍 Search Again"):
            try:
                with st.spinner("Searching again..."):
                    prompt = build_prompt(
                        st.session_state.research_input.get("topic", ""),
                        st.session_state.research_input.get("side", "").lower(),
                        st.session_state.research_input.get("claim", ""),
                        st.session_state.research_input.get("quant", False)
                    )
                    raw_result = query_perplexity(prompt)
                    st.session_state.evidence_result = raw_result
                    log_query(username, TOOL_NAME, True)
                    st.success("✅ Evidence generated!")
                    st.rerun()
            except Exception as e:
                st.error(f"❌ Error fetching evidence: {e}")
                log_query(username, TOOL_NAME, False)
    with col3:
        if st.button("🆕 New Search"):
            st.session_state.evidence_result = None
            st.session_state.research_input = {}
            st.rerun()

    # --- Feedback ---
    st.write("Was this useful?")
    feedback_col1, feedback_col2 = st.columns([1, 1])  # Same row, close together
    with feedback_col1:
        if st.button("👍 Yes"):
            log_query(username, TOOL_NAME, True, feedback="up")
            st.success("Thanks for the feedback!")
    with feedback_col2:
        if st.button("👎 No"):
            log_query(username, TOOL_NAME, True, feedback="down")
            st.info("We'll try to do better!")